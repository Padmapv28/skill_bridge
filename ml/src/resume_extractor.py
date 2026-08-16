"""
resume_extractor.py

Extracts raw plain text from resume files (PDF or DOCX).

Strategy:
    - PDF: pdfplumber (primary) -> PyMuPDF/fitz (fallback if pdfplumber
      returns little/no text, e.g. malformed PDFs or certain layouts)
    - DOCX: python-docx

Usage:
    from resume_extractor import extract_text
    text = extract_text("path/to/resume.pdf")
"""

import os
from pathlib import Path

import pdfplumber
import docx  # python-docx

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False


class ResumeExtractionError(Exception):
    """Raised when text cannot be extracted from a resume file."""
    pass


MIN_CHARS_PER_PAGE_THRESHOLD = 20


def _extract_pdf_with_pdfplumber(file_path: str) -> str:
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        if len(pdf.pages) == 0:
            raise ResumeExtractionError(f"PDF has no pages: {file_path}")
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_pdf_with_pymupdf(file_path: str) -> str:
    if not PYMUPDF_AVAILABLE:
        raise ResumeExtractionError(
            "PyMuPDF fallback requested but 'fitz' is not installed. "
            "Install it with: pip install PyMuPDF"
        )
    text_parts = []
    doc = fitz.open(file_path)
    try:
        if doc.page_count == 0:
            raise ResumeExtractionError(f"PDF has no pages: {file_path}")
        for page in doc:
            text_parts.append(page.get_text())
    finally:
        doc.close()
    return "\n".join(text_parts)


def _looks_like_scanned_or_empty(text: str, num_pages: int) -> bool:
    if num_pages == 0:
        return True
    avg_chars_per_page = len(text.strip()) / num_pages
    return avg_chars_per_page < MIN_CHARS_PER_PAGE_THRESHOLD


def _extract_from_pdf(file_path: str) -> str:
    try:
        with pdfplumber.open(file_path) as pdf:
            num_pages = len(pdf.pages)
    except Exception as e:
        raise ResumeExtractionError(f"Could not open PDF (corrupted?): {file_path} ({e})")

    try:
        text = _extract_pdf_with_pdfplumber(file_path)
    except ResumeExtractionError:
        raise
    except Exception:
        text = ""

    if text.strip() and not _looks_like_scanned_or_empty(text, num_pages):
        return text

    if PYMUPDF_AVAILABLE:
        try:
            fallback_text = _extract_pdf_with_pymupdf(file_path)
            if fallback_text.strip() and not _looks_like_scanned_or_empty(fallback_text, num_pages):
                return fallback_text
        except ResumeExtractionError:
            pass

    raise ResumeExtractionError(
        f"Could not extract meaningful text from '{file_path}'. "
        "This usually means the PDF is scanned/image-only and needs OCR, "
        "or the file is corrupted."
    )


def _extract_from_docx(file_path: str) -> str:
    try:
        document = docx.Document(file_path)
    except Exception as e:
        raise ResumeExtractionError(f"Could not open DOCX (corrupted?): {file_path} ({e})")

    text_parts = []

    for para in document.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for line in cell.text.split("\n"):
                    if line.strip():
                        text_parts.append(line.strip())

    full_text = "\n".join(text_parts)

    if not full_text.strip():
        raise ResumeExtractionError(
            f"DOCX file appears to have no extractable text: {file_path}"
        )

    return full_text


def extract_text(file_path: str) -> str:
    path = Path(file_path)

    if not path.exists():
        raise ResumeExtractionError(f"File not found: {file_path}")

    ext = path.suffix.lower()

    if ext == ".pdf":
        return _extract_from_pdf(str(path))
    elif ext == ".docx":
        return _extract_from_docx(str(path))
    else:
        raise ResumeExtractionError(
            f"Unsupported file type '{ext}'. Only .pdf and .docx are supported."
        )


if __name__ == "__main__":
    import sys
    import json

    if len(sys.argv) < 2:
        print("Usage: python resume_extractor.py <path_to_resume> [more paths...]")
        sys.exit(1)

    results = {}
    for fp in sys.argv[1:]:
        try:
            extracted = extract_text(fp)
            results[fp] = {
                "status": "success",
                "char_count": len(extracted),
                "preview": extracted[:200],
            }
        except ResumeExtractionError as e:
            results[fp] = {"status": "error", "message": str(e)}

    print(json.dumps(results, indent=2))
